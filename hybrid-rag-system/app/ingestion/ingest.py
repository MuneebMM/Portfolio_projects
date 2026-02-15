"""Document ingestion: loading and chunking text files."""

import json
import os
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.core.config import settings
from app.core.logger import logger


class DocumentLoader:
    """Loads .txt, .md, .pdf, .json, .docx, .doc, .html, .htm files from the data directory."""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".json", ".docx", ".doc", ".html", ".htm"}

    def __init__(self, data_dir: str | None = None):
        self.data_dir = Path(data_dir or settings.data_dir)

    # --- per-format parsers ---

    def _load_txt(self, file_path: Path) -> str:
        return file_path.read_text(encoding="utf-8")

    def _load_pdf(self, file_path: Path) -> str:
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages)

    def _load_docx(self, file_path: Path) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _load_html(self, file_path: Path) -> str:
        from bs4 import BeautifulSoup
        html = file_path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html, "lxml")
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    def _load_json(self, file_path: Path) -> str:
        data = json.loads(file_path.read_text(encoding="utf-8"))
        return json.dumps(data, indent=2, ensure_ascii=False)

    def _parse(self, file_path: Path) -> str:
        ext = file_path.suffix.lower()
        if ext in {".txt", ".md"}:
            return self._load_txt(file_path)
        if ext == ".pdf":
            return self._load_pdf(file_path)
        if ext in {".docx", ".doc"}:
            return self._load_docx(file_path)
        if ext in {".html", ".htm"}:
            return self._load_html(file_path)
        if ext == ".json":
            return self._load_json(file_path)
        raise ValueError(f"Unsupported extension: {ext}")

    # --- public API ---

    def load(self) -> list[Document]:
        documents = []
        if not self.data_dir.exists():
            logger.warning(f"Data directory not found: {self.data_dir}")
            return documents

        for file_path in sorted(self.data_dir.iterdir()):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    text = self._parse(file_path)
                    doc = Document(
                        page_content=text,
                        metadata={"source": file_path.name},
                    )
                    documents.append(doc)
                    logger.info(f"Loaded: {file_path.name} ({len(text)} chars)")
                except Exception as e:
                    logger.error(f"Failed to load {file_path.name}: {e}")

        logger.info(f"Total documents loaded: {len(documents)}")
        return documents


class TextChunker:
    """Splits documents into overlapping chunks."""

    def __init__(
        self,
        chunk_size: int | None = None,
        chunk_overlap: int | None = None,
    ):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size or settings.chunk_size,
            chunk_overlap=chunk_overlap or settings.chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len,
        )

    def chunk(self, documents: list[Document]) -> list[Document]:
        chunks = self.splitter.split_documents(documents)
        logger.info(f"Created {len(chunks)} chunks from {len(documents)} documents")
        return chunks


def ingest_documents(data_dir: str | None = None) -> list[Document]:
    """Full ingestion pipeline: load → chunk."""
    loader = DocumentLoader(data_dir)
    documents = loader.load()

    if not documents:
        logger.warning("No documents found to ingest")
        return []

    chunker = TextChunker()
    chunks = chunker.chunk(documents)
    return chunks
